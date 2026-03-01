from docx import Document
import sys
import os


def replace_placeholder(paragraph, placeholder, value):
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)


def generate_rbt(minggu, kelas, tarikh, hari, masa, refleksi):

    template_path = f"rbt_templates/M{minggu}.docx"

    if not os.path.exists(template_path):
        raise Exception("Template minggu tidak dijumpai.")

    doc = Document(template_path)

    # ==============================
    # Replace placeholder (paragraph)
    # ==============================
    for p in doc.paragraphs:
        replace_placeholder(p, "{{KELAS}}", kelas)
        replace_placeholder(p, "{{TARIKH}}", tarikh)
        replace_placeholder(p, "{{HARI}}", hari)
        replace_placeholder(p, "{{MASA}}", masa)
        replace_placeholder(p, "{{REFLEKSI}}", refleksi)

    # ==============================
    # Replace dalam TABLE
    # ==============================
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder(paragraph, "{{KELAS}}", kelas)
                    replace_placeholder(paragraph, "{{TARIKH}}", tarikh)
                    replace_placeholder(paragraph, "{{HARI}}", hari)
                    replace_placeholder(paragraph, "{{MASA}}", masa)
                    replace_placeholder(paragraph, "{{REFLEKSI}}", refleksi)

    # ==============================
    # Save output
    # ==============================
    output_folder = "output"
    os.makedirs(output_folder, exist_ok=True)

    output_file = f"{output_folder}/RPH_RBT_M{minggu}.docx"
    doc.save(output_file)

    return output_file


if __name__ == "__main__":

    minggu = sys.argv[1]
    kelas = sys.argv[2]
    tarikh = sys.argv[3]
    hari = sys.argv[4]
    masa = sys.argv[5]
    refleksi = sys.argv[6]

    file_path = generate_rbt(minggu, kelas, tarikh, hari, masa, refleksi)

    print("SIAP:", file_path)
